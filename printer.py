import time
from docx import Document

def getCurrentTime():
    return time.strftime("%d/%m/%Y | %H:%M:%S", time.localtime())

def ink(name='', age=-1, doctor='', indication='', medication='', tensAvgSist=-1.0, tensAvgDias=-1.0, tensMaxDaySist=-1.0, tensMaxDayDias=-1.0, tensMaxNightSist=-1.0, tensMaxNightDias=-1.0, variation=-1.0, sintoms=False, filepath=''):

	if name == '' or age == -1 or doctor == '' or indication == '' or medication == '' or tensAvgSist == -1.0 or tensAvgDias == -1.0 or tensMaxDaySist == -1.0 or tensMaxDayDias == -1.0 or tensMaxNightSist == -1.0 or tensMaxNightDias == -1.0 or variation == -1.0 or filepath == '':
		print('ERRO: Impossível imprimir relatório!')
		return False

	tensAvg = ''
	if tensAvgSist <= 139 and tensAvgDias <= 89:
		tensAvg = 'normais'
	else:
		tensAvg = 'elevados'

	tensMaxDay = ''
	if tensMaxDaySist <= 139 and tensMaxDayDias <= 89:
		tensMaxDay = 'normais'
	elif 140 <= tensMaxDaySist <= 159 and 90 <= tensMaxDayDias <= 99:
		tensMaxDay = 'de grau I'
	elif 160 <= tensMaxDaySist <= 179 and 100 <= tensMaxDayDias <= 109:
		tensMaxDay = 'de grau II'
	elif tensMaxDaySist >= 180 and tensMaxDayDias >= 110:
		tensMaxDay = 'de grau III'

	tensMaxNight = ''
	if tensMaxNightSist <= 139 and tensMaxNightDias <= 89:
		tensMaxNight = 'normais'
	elif 140 <= tensMaxNightSist <= 159 and 90 <= tensMaxNightDias <= 99:
		tensMaxNight = 'de grau I'
	elif 160 <= tensMaxNightSist <= 179 and 100 <= tensMaxNightDias <= 109:
		tensMaxNight = 'de grau II'
	elif tensMaxNightSist >= 180 and tensMaxNightDias >= 110:
		tensMaxNight = 'de grau III'

	behav = 'dipper' if variation > 10 else 'não dipper'

	sintom = 'Presença' if sintoms else 'Ausência'

	document = Document()

	document.add_heading('Relatório médico\n', 1)

	p = document.add_paragraph()
	p.add_run(name).bold = True
	p.add_run(' , ')
	p.add_run(str(age)).bold = True
	p.add_run(' anos')

	p = document.add_paragraph('Dr./Dra. ')
	p.add_run(doctor).bold = True

	p = document.add_paragraph('Indicação: ')
	p.add_run(indication).bold = True
	p.add_run('.\n')

	document.add_paragraph('Relatório:')

	p = document.add_paragraph('Paciente medicado com ')
	p.add_run(medication).bold = True
	p.add_run('.')

	p = document.add_paragraph('Valores tensionais médios ')
	p.add_run(tensAvg).bold = True
	p.add_run(', segundo o cálculo das cargas tensionais.')

	p = document.add_paragraph('Registo de valores tensionais ')
	p.add_run(tensMaxDay).bold = True
	p.add_run(' no período diurno.')

	p = document.add_paragraph('Registo de valores tensionais ')
	p.add_run(tensMaxNight).bold = True
	p.add_run(' no período noturno.')

	p = document.add_paragraph('Comportamento ')
	p.add_run(behav).bold = True
	p.add_run('.')

	p = document.add_paragraph()
	p.add_run(sintom).bold = True
	p.add_run(' de sintomas durante o registo.\n')

	document.add_paragraph('Relatório gerado em ' + getCurrentTime())

	open(filepath, 'w')
	document.save(filepath)

	return True

